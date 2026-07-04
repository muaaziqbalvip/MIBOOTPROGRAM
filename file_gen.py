"""
MI AI - File Generation
PDF, DOCX, aur full website (HTML/CSS/JS) ZIP me generate karta hai
"""

import os
import time
import zipfile
from docx import Document
from docx.shared import Pt
from fpdf import FPDF

from config import TEMP_DIR


def generate_pdf(title: str, content: str) -> str:
    pdf = FPDF()
    pdf.add_page()

    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
    if os.path.exists(font_path):
        pdf.add_font("DejaVu", "", font_path, uni=True)
        pdf.set_font("DejaVu", size=16)
    else:
        pdf.set_font("Arial", size=16)

    pdf.multi_cell(0, 10, title)
    pdf.ln(5)

    if os.path.exists(font_path):
        pdf.set_font("DejaVu", size=12)
    else:
        pdf.set_font("Arial", size=12)

    pdf.multi_cell(0, 8, content)

    filename = f"document_{int(time.time())}.pdf"
    filepath = os.path.join(TEMP_DIR, filename)
    pdf.output(filepath)
    return filepath


def generate_docx(title: str, content: str) -> str:
    doc = Document()
    doc.add_heading(title, level=1)

    for paragraph in content.split("\n"):
        if paragraph.strip():
            p = doc.add_paragraph(paragraph)
            p.style.font.size = Pt(12)

    filename = f"document_{int(time.time())}.docx"
    filepath = os.path.join(TEMP_DIR, filename)
    doc.save(filepath)
    return filepath


def generate_website_zip(site_name: str, html: str, css: str, js: str) -> str:
    folder_name = f"website_{int(time.time())}"
    folder_path = os.path.join(TEMP_DIR, folder_name)
    os.makedirs(folder_path, exist_ok=True)

    with open(os.path.join(folder_path, "index.html"), "w", encoding="utf-8") as f:
        f.write(html)
    with open(os.path.join(folder_path, "style.css"), "w", encoding="utf-8") as f:
        f.write(css)
    with open(os.path.join(folder_path, "script.js"), "w", encoding="utf-8") as f:
        f.write(js)

    zip_path = os.path.join(TEMP_DIR, f"{site_name}_{int(time.time())}.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
        for file in os.listdir(folder_path):
            zipf.write(os.path.join(folder_path, file), arcname=file)

    return zip_path


def files_to_zip(zip_name: str, files: dict) -> str:
    """files = { "filename.txt": "content string", ... }"""
    zip_path = os.path.join(TEMP_DIR, f"{zip_name}_{int(time.time())}.zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zipf:
        for filename, content in files.items():
            zipf.writestr(filename, content)
    return zip_path